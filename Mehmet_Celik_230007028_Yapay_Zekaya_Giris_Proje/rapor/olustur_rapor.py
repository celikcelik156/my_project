import os
import pickle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MODEL_DOSYASI = os.path.join(PROJE_DIR, "data", "modeller.pkl")
GRAFIK_DIR = os.path.join(PROJE_DIR, "grafikler")
RAPOR_DOSYASI = os.path.join(PROJE_DIR, "rapor", "teknik_rapor.docx")


def baslik(doc, text, level=1):
    doc.add_heading(text, level=level)


def paragraf(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def olustur():
    with open(MODEL_DOSYASI, "rb") as f:
        kayit = pickle.load(f)
    metrikler = kayit["reg_metrikler"]
    en_iyi = max(metrikler, key=lambda m: m["R2"])

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    t = doc.add_heading("DONEM SONU PROJESI TEKNIK RAPORU", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraf(doc, "Ders: Yapay Zekaya Giris")
    paragraf(doc, "Konu: Turkiye Emlak Fiyat Tahmin Sistemi")
    paragraf(doc, "Ogrenci: Mehmet Celik")
    paragraf(doc, "Ogrenci No: 230007028")
    doc.add_page_break()

    baslik(doc, "1. Giris")
    paragraf(doc, "Bu proje, konut ozelliklerine gore satis fiyati tahmini yapan basit bir makine ogrenmesi uygulamasidir. Gercek hayatta emlak alim-satim kararlarinda fiyat tahmini onemli bir problemdir. Proje Python ve MATLAB ortamlarinda gelistirilmis, Streamlit ile kullanici arayuzu sunulmustur.")
    paragraf(doc, "Proje bilincli olarak sade tutulmustur: veri on isleme (normalizasyon, encoding, ozellik muhendisligi) uygulanmamistir. Yalnizca CSV dosyasindaki ham sayisal ozellikler dogrudan modele verilmistir.")

    baslik(doc, "2. Veri Seti Nedir?")
    paragraf(doc, "Veri seti, her satiri bir konutu temsil eden 1200 kayittan olusan tablo yapisi (CSV) formatindaki veridir. Veri setinde fiyat (hedef degisken) ile birlikte metrekare, oda sayisi, bina yasi, bulundugu kat, toplam kat, balkon, esyali, site icinde ve ulasim skoru gibi sayisal alanlar bulunmaktadir.")
    paragraf(doc, "Bu calismada sehir, ilce, isinma gibi kategorik sutunlar modele dahil edilmemistir. Boylece veri on isleme ihtiyaci ortadan kaldirilmis, proje minimum duzeyde tutulmustur.")

    baslik(doc, "3. Model Nedir?")
    paragraf(doc, "Model, girdiler (ozellikler) ile cikti (fiyat) arasindaki iliskiyi ogrenen matematiksel bir yapidir. Regresyon modelleri surekli bir deger tahmin eder. Bu projede 5 regresyon algoritmasi kullanilmistir:")
    for m in ["Linear Regression", "Decision Tree Regressor", "Random Forest Regressor", "KNN Regressor", "SVM Regressor"]:
        doc.add_paragraph(m, style="List Bullet")

    baslik(doc, "4. Model Nasil Egitilir?")
    paragraf(doc, "Egitim sureci su adimlarla yapilmistir:")
    doc.add_paragraph("CSV dosyasi okunur (1200 kayit).", style="List Number")
    doc.add_paragraph("Veri %80 egitim, %20 test olarak ayrilir (random_state=42).", style="List Number")
    doc.add_paragraph("Her algoritma egitim verisi ile fit() metodu kullanilarak egitilir.", style="List Number")
    doc.add_paragraph("Egitilen modeller kaydedilir (modeller.pkl).", style="List Number")
    paragraf(doc, "Model egitimi sirasinda test verisi kullanilmaz; model yalnizca egitim verisindeki oruntuleri ogrenir.")

    baslik(doc, "5. Model Nasil Test Edilir?")
    paragraf(doc, "Egitim tamamlandiktan sonra model, daha once gormedigi test verisi uzerinde predict() ile degerlendirilir. Gercek fiyat ile tahmin karsilastirilarak MAE, MSE, RMSE ve R2 metrikleri hesaplanir.")
    paragraf(doc, "MAE ortalama mutlak hatayi, RMSE buyuk hatalari daha cok cezalandirir, R2 ise modelin aciklama gucunu gosterir (1'e yakin daha iyi).")

    baslik(doc, "6. Python Sonuclari (Veri On Isleme Yok)")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Model", "MAE (TL)", "RMSE (TL)", "R2", "Yorum"]):
        hdr[i].text = h
    for m in metrikler:
        row = table.add_row().cells
        row[0].text = m["model"]
        row[1].text = f"{m['MAE']:,.0f}"
        row[2].text = f"{m['RMSE']:,.0f}"
        row[3].text = f"{m['R2']:.4f}"
        if m["R2"] == en_iyi["R2"]:
            row[4].text = "En iyi R2"
        elif m["R2"] < 0:
            row[4].text = "On isleme olmadan zayif"
        else:
            row[4].text = "-"
    paragraf(doc, f"En iyi model: {en_iyi['model']} (R2 = {en_iyi['R2']:.4f}). SVM, olcek farklari nedeniyle on isleme olmadan dusuk performans gostermistir.")

    baslik(doc, "7. MATLAB Uygulamasi")
    paragraf(doc, "Ayni veri seti ve ayni sayisal ozellikler MATLAB ortaminda da kullanilmistir. emlak_analiz.m dosyasi fitlm, fitrtree, TreeBagger, knnsearch ve fitrsvm fonksiyonlari ile 5 algoritmayi calistirir. Veri on isleme yapilmamistir.")

    baslik(doc, "8. Kullanici Arayuzu")
    paragraf(doc, "Streamlit tabanli web arayuzu (ui/app.py) ile kullanici konut ozelliklerini girerek 5 modelden tahmin alabilir ve model karsilastirma tablosunu gorebilir.")

    baslik(doc, "9. Grafiksel Analizler")
    paragraf(doc, "Asagidaki grafikler main.py calistirildiginda otomatik uretilmistir:")
    for g in ["01_algoritma_karsilastirma.png", "02_gercek_tahmin_random_forest.png", "03_korelasyon.png"]:
        yol = os.path.join(GRAFIK_DIR, g)
        if os.path.exists(yol):
            doc.add_paragraph(g.replace("_", " ").replace(".png", ""))
            doc.add_picture(yol, width=Inches(5.5))

    baslik(doc, "10. Sonuc")
    paragraf(doc, "Bu projede veri on isleme yapilmadan 5 regresyon algoritmasi karsilastirilmistir. Linear Regression en yuksek R2 degerini vermistir. Proje, ders kapsaminda makine ogrenmesi kavramlarini (veri seti, model, egitim, test) basit ve anlasilir bir uygulama ile gostermektedir.")
    paragraf(doc, "Gelistirici: Mehmet Celik | 230007028 | Yapay Zekaya Giris")

    os.makedirs(os.path.dirname(RAPOR_DOSYASI), exist_ok=True)
    doc.save(RAPOR_DOSYASI)
    print(f"Rapor olusturuldu: {RAPOR_DOSYASI}")


if __name__ == "__main__":
    olustur()
